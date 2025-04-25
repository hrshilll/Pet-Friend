from django.contrib import admin
from django.contrib.auth.admin import UserAdmin
from django.http import HttpResponse
from .models import CustomUser, Product, Service
from docx import Document
import os
from django.conf import settings
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== User Export =====
def export_users_docx(modeladmin, request, queryset):
    root_path = settings.BASE_DIR
    save_path = os.path.join(root_path, 'users.docx')

    doc = Document()
    title = doc.add_heading('User List', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=10)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    hdr_cells = table.rows[0].cells
    headers = ['Username', 'Email', 'Full Name', 'Phone Number', 'Address',
               'Pet Name', 'Pet Type', 'Pet Breed', 'Pet Gender', 'Pet Age']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(11)

    for user in queryset:
        row = table.add_row().cells
        row[0].text = user.username
        row[1].text = user.email
        row[2].text = str(user.full_name)
        row[3].text = str(user.phone_number)
        row[4].text = str(user.address)
        row[5].text = str(user.pet_name)
        row[6].text = str(user.pet_type)
        row[7].text = str(user.pet_breed)
        row[8].text = str(user.pet_gender)
        row[9].text = str(user.pet_age)

    doc.save(save_path)

    with open(save_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=users.docx'
        return response

export_users_docx.short_description = "Print selected users as DOCX"

# ===== Product Export =====
def export_products_docx(modeladmin, request, queryset):
    root_path = settings.BASE_DIR
    save_path = os.path.join(root_path, 'products.docx')

    doc = Document()
    title = doc.add_heading('Product List', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Price'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(11)

    for product in queryset:
        row = table.add_row().cells
        row[0].text = product.name
        row[1].text = str(product.price)

    doc.save(save_path)

    with open(save_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=products.docx'
        return response

export_products_docx.short_description = "Print selected products as DOCX"

def export_services_docx(modeladmin, request, queryset):
    root_path = settings.BASE_DIR
    save_path = os.path.join(root_path, 'services.docx')

    doc = Document()
    title = doc.add_heading('Service List', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Service Name'
    hdr_cells[1].text = 'Price'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0]
            run.bold = True
            run.font.size = Pt(11)

    for service in queryset:
        row = table.add_row().cells
        row[0].text = service.name
        row[1].text = str(service.price)

    doc.save(save_path)

    with open(save_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=services.docx'
        return response

export_services_docx.short_description = "Print selected services as DOCX"
# ===== Admin Registrations =====

@admin.register(CustomUser)
class CustomUserAdmin(UserAdmin):
    model = CustomUser
    actions = [export_users_docx]
    fieldsets = UserAdmin.fieldsets + (
        ('Additional Info', {
            'fields': (
                'full_name', 'phone_number', 'address',
                'pet_name', 'pet_type', 'pet_breed',
                'pet_age', 'pet_gender'
            )
        }),
    )

@admin.register(Product)
class ProductAdmin(admin.ModelAdmin):
    list_display = ('name', 'price', 'image' ,"stock")
    list_editable = ('price', 'image', 'stock')
    list_display_links = ('name',)
    actions = ['delete_selected', export_products_docx]

    def get_queryset(self, request):
        qs = super().get_queryset(request)
        return qs.order_by('stock')

@admin.register(Service)
class ServiceAdmin(admin.ModelAdmin):
    list_display = ('name', 'price', 'category', 'slot')  # Show in list view
    list_filter = ('category', 'slot')  # Enable filter sidebar
    search_fields = ('name', 'category')  # Optional: improve search
    actions = [export_services_docx]